"""Convert report.md to formatted .docx using python-docx."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

INPUT = r'C:\Users\wucan\IdeaProjects\movie-database\report.md'
OUTPUT = r'C:\Users\wucan\Desktop\CineBase实验报告.docx'

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

# Styles
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

def add_styled_paragraph(text, bold=False, size=12, align=None, font_name='宋体'):
    """Add a paragraph with the given styling."""
    # Remove markdown formatting chars for docx
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # bold markers
    text = re.sub(r'__(.+?)__', r'\1', text)
    para = doc.add_paragraph()
    if align:
        para.alignment = align
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold
    return para

def add_heading_styled(text, level):
    """Add a heading."""
    sizes = {1: 22, 2: 16, 3: 14}
    size = sizes.get(level, 14)
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Arial'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    return para

# Read markdown file
with open(INPUT, 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip()
    i += 1

    # Skip empty lines and frontmatter
    if not line or line.startswith('---') or line.startswith('documentclass'):
        continue

    # Headings
    m = re.match(r'^(#{1,3})\s+(.+)$', line)
    if m:
        level = len(m.group(1))
        text = m.group(2)
        # Skip markdown heading markers
        text = re.sub(r'\\(.*?)\{.*?\}', r'\1', text)  # LaTeX commands
        add_heading_styled(text, level)
        continue

    # Table rows
    if line.startswith('|'):
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if all(c.startswith('---') or c.startswith(':--') for c in cells if c):
            continue  # skip separator row
        if any(cells):
            para = doc.add_paragraph()
            text = ' | '.join(c for c in cells if c)
            run = para.add_run(text)
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        continue

    # Horizontal rule (page break)
    if line == r'\newpage':
        doc.add_page_break()
        continue

    # Regular paragraph
    # Clean LaTeX from the text
    cleaned = re.sub(r'\\(?:textbf|textit|texttt)\{(.+?)\}', r'\1', line)
    cleaned = re.sub(r'\\newpage', '', cleaned)
    cleaned = cleaned.strip()

    if cleaned:
        # Handle bold markers **text**
        parts = re.split(r'(\*\*.+?\*\*)', cleaned)
        para = doc.add_paragraph()
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = 'Arial'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            else:
                run = para.add_run(part)
                run.font.size = Pt(12)
                run.font.name = 'Arial'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.save(OUTPUT)
print(f'Report generated: {OUTPUT}')
print('Done!')
